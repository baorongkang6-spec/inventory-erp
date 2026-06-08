#!/usr/bin/env python3
"""操作手册 Markdown → 带真表格的 Word(.docx)。中文字体用宋体/黑体。"""
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

src, out = sys.argv[1], sys.argv[2]
lines = open(src, encoding="utf-8").read().splitlines()

doc = Document()
# 默认中文字体
normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_runs(p, text):
    """解析 **加粗** 和 `代码`，加到段落 p。"""
    for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xC0, 0x26, 0xD3)
        else:
            p.add_run(part)


i, n = 0, len(lines)
while i < n:
    s = lines[i].strip()
    if not s:
        i += 1; continue
    # 表格
    if s.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?\s*$", lines[i + 1].strip()):
        header = [c.strip() for c in s.strip("|").split("|")]
        i += 2; rows = []
        while i < n and lines[i].strip().startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
        t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
        for j, c in enumerate(header):
            cell = t.rows[0].cells[j]; cell.paragraphs[0].clear()
            add_runs(cell.paragraphs[0], c)
            for run in cell.paragraphs[0].runs: run.bold = True
        for r in rows:
            cells = t.add_row().cells
            for j, c in enumerate(r):
                if j < len(cells):
                    cells[j].paragraphs[0].clear(); add_runs(cells[j].paragraphs[0], c)
        doc.add_paragraph(); continue
    # 标题
    m = re.match(r"^(#{1,6})\s+(.*)$", s)
    if m:
        doc.add_heading(m.group(2), level=min(len(m.group(1)) - 1, 4)); i += 1; continue
    # 分割线
    if re.match(r"^---+$", s):
        doc.add_paragraph("─" * 30); i += 1; continue
    # 引用
    if s.startswith(">"):
        q = []
        while i < n and lines[i].strip().startswith(">"):
            q.append(lines[i].strip()[1:].strip()); i += 1
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
        add_runs(p, "  ".join(q))
        for run in p.runs: run.italic = True; run.font.color.rgb = RGBColor(0x74, 0x42, 0x10)
        continue
    # 有序列表
    if re.match(r"^\d+\.\s+", s):
        while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
            p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s+", "", lines[i].strip())); i += 1
        continue
    # 无序列表
    if s.startswith("- "):
        while i < n and lines[i].strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, lines[i].strip()[2:]); i += 1
        continue
    # 普通段落
    p = doc.add_paragraph(); add_runs(p, s); i += 1

doc.save(out)
print("docx 生成:", out, " 表格数:", len(doc.tables))
